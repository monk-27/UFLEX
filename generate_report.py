#!/usr/bin/env python3
"""Generate a professional VAPT Remediation Report DOCX for UFlex Ltd."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ─── Colors ───
DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
ACCENT_BLUE = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BLUE_BG = "D6E4F0"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
RED_ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREEN_ACCENT = RGBColor(0x27, 0xAE, 0x60)
ORANGE_ACCENT = RGBColor(0xE6, 0x7E, 0x22)

doc = Document()

# ─── Default font ───
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = DARK_GRAY

# ─── Heading styles ───
for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.color.rgb = DARK_BLUE
    h.font.bold = True
    pf = h.paragraph_format
    pf.space_before = Pt(18 if i == 1 else 14)
    pf.space_after = Pt(8)

doc.styles['Heading 1'].font.size = Pt(22)
doc.styles['Heading 2'].font.size = Pt(16)
doc.styles['Heading 3'].font.size = Pt(13)

# ─── Page margins ───
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(headers, rows, col_widths=None, header_color="1B3A5C"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, header_color)

    # Data rows
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
            run.font.color.rgb = DARK_GRAY
            if i % 2 == 1:
                set_cell_shading(cell, "F2F7FB")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing
    return table


def add_info_box(text, box_color=LIGHT_BLUE_BG):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_BLUE
    run.italic = True
    # Add background shading to paragraph
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{box_color}" w:val="clear"/>')
    pPr.append(shading)


def add_code_block(code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shading)


def add_severity_badge(text, color_label="low"):
    colors = {
        "low": ("FFF3CD", "856404"),
        "info": ("D1ECF1", "0C5460"),
        "medium": ("FFE0CC", "CC5500"),
        "high": ("F8D7DA", "721C24"),
    }
    bg, fg = colors.get(color_label, colors["low"])
    p = doc.add_paragraph()
    run = p.add_run(f"  Severity: {text}  ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor(int(fg[:2], 16), int(fg[2:4], 16), int(fg[4:], 16))
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg}" w:val="clear"/>')
    pPr.append(shading)


def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK_GRAY
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GRAY


def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════

# Add spacing at top
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("UFLEX LTD")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = DARK_BLUE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Security Enhancement Report")
run.font.size = Pt(20)
run.font.color.rgb = ACCENT_BLUE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("VAPT Findings & Remediation Plan")
run.font.size = Pt(14)
run.font.color.rgb = MEDIUM_GRAY
run.font.name = 'Calibri'

doc.add_paragraph()
doc.add_paragraph()

# Info box on cover
cover_table = doc.add_table(rows=4, cols=2)
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cover_info = [
    ("Prepared For", "UFlex Ltd"),
    ("Date", "April 16, 2026"),
    ("Scope", "beta.uflexltd.com  |  uflexltd.com"),
    ("Server", "184.168.125.211"),
]
for i, (label, value) in enumerate(cover_info):
    cell_l = cover_table.rows[i].cells[0]
    cell_r = cover_table.rows[i].cells[1]
    cell_l.text = ""
    cell_r.text = ""
    run_l = cell_l.paragraphs[0].add_run(label)
    run_l.bold = True
    run_l.font.size = Pt(11)
    run_l.font.name = 'Calibri'
    run_l.font.color.rgb = DARK_BLUE
    run_r = cell_r.paragraphs[0].add_run(value)
    run_r.font.size = Pt(11)
    run_r.font.name = 'Calibri'
    run_r.font.color.rgb = DARK_GRAY
    set_cell_shading(cell_l, "F2F7FB")
    cell_l.width = Inches(2)
    cell_r.width = Inches(4)

# Page break
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════

doc.add_heading("Table of Contents", level=1)

toc_items = [
    "1.  Executive Summary",
    "2.  What's Changing in the Beta Version",
    "",
    "SECTION A: beta.uflexltd.com (New Beta Website)",
    "    A1.  Adding Content Security Policy (CSP)",
    "    A2.  Adding X-Content-Type-Options Header",
    "    A3.  Adding Referrer-Policy Header",
    "    A4.  Adding Permissions-Policy Header",
    "    A5.  Adding Cache-Control Headers",
    "    A6.  Infrastructure Hardening",
    "    Complete .htaccess Configuration",
    "",
    "SECTION B: uflexltd.com (Current Production Website)",
    "    B1-B5.  Security Headers",
    "    B6.  Infrastructure Hardening",
    "    B7.  Privacy & Legal Pages",
    "",
    "SECTION C: Server-Level Updates",
    "    C1.  Exim Mail Server Update",
    "    C2.  OpenSSH Update",
    "",
    "Remediation Roadmap",
    "Summary of Impact",
]

for item in toc_items:
    if item == "":
        doc.add_paragraph()
        continue
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    if item.startswith("SECTION") or item == "Remediation Roadmap" or item == "Summary of Impact":
        run.bold = True
        run.font.color.rgb = DARK_BLUE
    elif item.startswith("    "):
        run.font.color.rgb = MEDIUM_GRAY
    else:
        run.font.color.rgb = DARK_GRAY

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════

doc.add_heading("1. Executive Summary", level=1)

p = doc.add_paragraph()
run = p.add_run(
    "Following a routine security assessment (VAPT) conducted on April 13, 2026, "
    "we have identified several areas where the security posture of both "
)
run.font.name = 'Calibri'
run.font.size = Pt(10.5)
run = p.add_run("beta.uflexltd.com")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(10.5)
run = p.add_run(" and ")
run.font.name = 'Calibri'
run.font.size = Pt(10.5)
run = p.add_run("uflexltd.com")
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(10.5)
run = p.add_run(
    " can be strengthened. Most of these are standard security headers that are "
    "straightforward to implement, along with two server-level software updates "
    "that the hosting provider can handle."
)
run.font.name = 'Calibri'
run.font.size = Pt(10.5)

add_info_box(
    "None of these findings indicate an active breach or data compromise. "
    "They represent opportunities to harden the infrastructure and align "
    "with industry best practices before the beta site goes live."
)

add_styled_table(
    ["Category", "Items", "Severity Level"],
    [
        ["Security Headers (Application Layer)", "5 improvements", "Low / Informational"],
        ["Infrastructure Hardening", "1 composite improvement", "Medium"],
        ["Server Software Updates", "2 updates needed", "High"],
    ],
    col_widths=[2.5, 2.0, 1.8]
)

add_divider()

# ═══════════════════════════════════════════════════════════════
# WHAT'S CHANGING
# ═══════════════════════════════════════════════════════════════

doc.add_heading("2. What's Changing in the Beta Version", level=1)

p = doc.add_paragraph(
    "The new beta.uflexltd.com is a complete rebuild of the website using modern technology. "
    "Here is a comparison of the current and new platforms:"
)

add_styled_table(
    ["Aspect", "Current Site (uflexltd.com)", "New Beta Site (beta.uflexltd.com)"],
    [
        ["Technology", "PHP (server-rendered pages)", "Next.js 15 + React 19 + TypeScript"],
        ["Styling", "Traditional CSS", "Tailwind CSS + Modern UI Components"],
        ["Performance", "Standard page loads", "Static pre-built pages (faster load times)"],
        ["Forms & APIs", "PHP form handlers", "Modern API routes + CAPTCHA validation"],
        ["Deployment", "Manual uploads", "Automated via GitHub Actions"],
        ["Database", "Direct PHP queries", "Connection pooling with MySQL"],
    ],
    col_widths=[1.5, 2.4, 2.4]
)

add_info_box(
    "Both sites currently share the same server (184.168.125.211), which means "
    "server-level improvements will benefit both environments simultaneously."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION A
# ═══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SECTION A")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = DARK_BLUE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("beta.uflexltd.com")
run.font.size = Pt(18)
run.font.color.rgb = ACCENT_BLUE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("New Beta Website - Next.js Application")
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_GRAY
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph(
    "The following improvements will be implemented as part of the beta site deployment. "
    "Since the beta site uses a static export (pre-built HTML/CSS/JS files served via Apache on cPanel), "
    "all security headers will be configured at the server level using .htaccess."
)

add_divider()

# ── A1: CSP ──
doc.add_heading("A1. Adding Content Security Policy (CSP)", level=2)
add_severity_badge("Low", "low")

p = doc.add_paragraph()
run = p.add_run("Reference: ")
run.bold = True
run.font.name = 'Calibri'
run = p.add_run("APPSEC-002")
run.font.name = 'Calibri'

doc.add_heading("In Simple Terms", level=3)
p = doc.add_paragraph(
    "A Content Security Policy is like a guest list for a website. It tells the browser exactly "
    "which external sources (scripts, fonts, images, analytics) are allowed to load on the page. "
    "Without this guest list, the browser allows anything and everything \u2014 which could include "
    "malicious scripts if someone tries to inject them."
)

doc.add_heading("What We Will Do", level=3)
p = doc.add_paragraph(
    "We will add a CSP header that explicitly whitelists only the trusted sources the beta site "
    "actually uses (Google Fonts, Google Analytics, Vercel Analytics). Everything else will be "
    "blocked by default."
)

doc.add_heading("Implementation", level=3)
add_code_block(
    'Header set Content-Security-Policy "default-src \'self\'; '
    'script-src \'self\' \'unsafe-inline\' \'unsafe-eval\' '
    'https://www.googletagmanager.com https://www.google-analytics.com '
    'https://va.vercel-scripts.com; '
    'style-src \'self\' \'unsafe-inline\' https://fonts.googleapis.com; '
    'font-src \'self\' https://fonts.gstatic.com; '
    'img-src \'self\' data: https: blob:; '
    'connect-src \'self\' https://www.google-analytics.com '
    'https://vitals.vercel-insights.com; '
    'frame-ancestors \'self\'; base-uri \'self\'; form-action \'self\';"'
)

add_divider()

# ── A2: X-Content-Type-Options ──
doc.add_heading("A2. Adding X-Content-Type-Options Header", level=2)
add_severity_badge("Low", "low")

p = doc.add_paragraph()
run = p.add_run("Reference: ")
run.bold = True
run = p.add_run("APPSEC-006")

doc.add_heading("In Simple Terms", level=3)
p = doc.add_paragraph(
    "When a browser receives a file, it sometimes tries to guess what type of file it is "
    "(called \"MIME sniffing\") instead of trusting what the server says. By adding this header, "
    "we instruct the browser to always trust the declared file type and never guess. "
    "This prevents scenarios where a mislabeled file could be executed as code."
)

doc.add_heading("What We Will Do", level=3)
p = doc.add_paragraph(
    "We will add a single header that prevents browsers from guessing file types. "
    "This is a simple, zero-risk addition."
)

doc.add_heading("Implementation", level=3)
add_code_block('Header set X-Content-Type-Options "nosniff"')

add_divider()

# ── A3: Referrer-Policy ──
doc.add_heading("A3. Adding Referrer-Policy Header", level=2)
add_severity_badge("Low", "low")

p = doc.add_paragraph()
run = p.add_run("Reference: ")
run.bold = True
run = p.add_run("APPSEC-003")

doc.add_heading("In Simple Terms", level=3)
p = doc.add_paragraph(
    "When a user clicks a link on the UFlex site that goes to an external website, "
    "the browser may send the full URL of the page they were on to that external site. "
    "A Referrer-Policy controls how much of this information is shared. We can configure it "
    "so that only the domain name (https://beta.uflexltd.com) is shared with external sites, "
    "keeping internal page paths private."
)

doc.add_heading("What We Will Do", level=3)
p = doc.add_paragraph(
    "We will set the policy to \"strict-origin-when-cross-origin\" \u2014 this shares the full URL "
    "only within the same site, and shares just the domain name when navigating to external sites."
)

doc.add_heading("Implementation", level=3)
add_code_block('Header set Referrer-Policy "strict-origin-when-cross-origin"')

add_divider()

# ── A4: Permissions-Policy ──
doc.add_heading("A4. Adding Permissions-Policy Header", level=2)
add_severity_badge("Low", "low")

p = doc.add_paragraph()
run = p.add_run("Reference: ")
run.bold = True
run = p.add_run("APPSEC-004")

doc.add_heading("In Simple Terms", level=3)
p = doc.add_paragraph(
    "Modern browsers have access to device features like camera, microphone, and location. "
    "A Permissions-Policy header lets us declare upfront that the website does not need access "
    "to these features. This prevents any embedded content (such as third-party widgets) "
    "from attempting to use them without authorization."
)

doc.add_heading("What We Will Do", level=3)
p = doc.add_paragraph(
    "We will disable all sensitive browser features that the UFlex website does not require \u2014 "
    "camera, microphone, geolocation, and payment APIs."
)

doc.add_heading("Implementation", level=3)
add_code_block('Header set Permissions-Policy "camera=(), microphone=(), geolocation=(), payment=(), fullscreen=(self)"')

add_divider()

# ── A5: Cache-Control ──
doc.add_heading("A5. Adding Cache-Control Headers", level=2)
add_severity_badge("Informational", "info")

p = doc.add_paragraph()
run = p.add_run("Reference: ")
run.bold = True
run = p.add_run("APPSEC-009")

doc.add_heading("In Simple Terms", level=3)
p = doc.add_paragraph(
    "Cache-Control tells browsers and intermediate servers how long they should store "
    "a copy of a page or file before checking for a new version. Without these instructions, "
    "caching behavior is unpredictable \u2014 users may see outdated content, or sensitive pages "
    "may be stored longer than intended."
)

doc.add_heading("What We Will Do", level=3)
add_bullet("will always check for the latest version (no stale content)", "HTML pages ")
add_bullet("(images, CSS, JS, fonts) will be cached for long periods since they are versioned during each build", "Static files ")
add_bullet("will not be cached at all", "API responses ")

doc.add_heading("Implementation", level=3)
add_code_block(
    '# HTML pages\n'
    '<FilesMatch "\\.(html)$">\n'
    '  Header set Cache-Control "no-cache, must-revalidate"\n'
    '</FilesMatch>\n\n'
    '# Static assets\n'
    '<FilesMatch "\\.(css|js|jpg|jpeg|png|gif|svg|woff|woff2)$">\n'
    '  Header set Cache-Control "public, max-age=31536000, immutable"\n'
    '</FilesMatch>'
)

add_divider()

# ── A6: Infrastructure Hardening ──
doc.add_heading("A6. Infrastructure Hardening for Beta", level=2)
add_severity_badge("Medium", "medium")

p = doc.add_paragraph()
run = p.add_run("Reference: ")
run.bold = True
run = p.add_run("ATTSRF-001")

add_styled_table(
    ["Improvement", "Description", "How"],
    [
        ["Rate Limiting on Forms", "Prevent automated bots from spamming the quote and contact forms", "Add request throttling to PHP form endpoints"],
        ["Hide Server Information", "Remove server software version details from HTTP responses", "Add ServerSignature Off and Header unset Server"],
        ["Web Application Firewall", "Add a protective layer that filters malicious traffic", "Deploy Cloudflare (free tier available) as a DNS-level WAF"],
        ["Additional Security Headers", "Add extra protective headers beyond the VAPT findings", "Add X-Frame-Options, X-XSS-Protection, and HSTS"],
    ],
    col_widths=[1.6, 2.2, 2.4]
)

add_divider()

# ── Complete .htaccess ──
doc.add_heading("Complete .htaccess for beta.uflexltd.com", level=2)

p = doc.add_paragraph(
    "All improvements A1 through A6 can be deployed in a single .htaccess file update:"
)

add_code_block(
    '# ============================================\n'
    '# UFlex Beta - Security Headers\n'
    '# Location: public_html/.htaccess (beta subdomain)\n'
    '# ============================================\n\n'
    '<IfModule mod_headers.c>\n'
    '  # A1: Content Security Policy\n'
    '  Header set Content-Security-Policy "default-src \'self\'; ..."\n\n'
    '  # A2: Prevent MIME Sniffing\n'
    '  Header set X-Content-Type-Options "nosniff"\n\n'
    '  # A3: Referrer Policy\n'
    '  Header set Referrer-Policy "strict-origin-when-cross-origin"\n\n'
    '  # A4: Permissions Policy\n'
    '  Header set Permissions-Policy "camera=(), microphone=(), ..."\n\n'
    '  # A5: Cache Control for HTML\n'
    '  Header set Cache-Control "no-cache, must-revalidate"\n\n'
    '  # A6: Additional Security Headers\n'
    '  Header set X-Frame-Options "SAMEORIGIN"\n'
    '  Header set X-XSS-Protection "1; mode=block"\n'
    '  Header set Strict-Transport-Security "max-age=31536000; ..."\n\n'
    '  # A6: Hide Server Information\n'
    '  Header unset X-Powered-By\n'
    '  Header unset Server\n'
    '</IfModule>\n\n'
    'ServerSignature Off'
)

add_info_box(
    "The full .htaccess with complete CSP and Permissions-Policy values "
    "is available in the accompanying technical reference file."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION B
# ═══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SECTION B")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = DARK_BLUE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("uflexltd.com")
run.font.size = Pt(18)
run.font.color.rgb = ACCENT_BLUE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Current Production Website - PHP Application")
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_GRAY
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph(
    "The same security headers and improvements apply to the production site. "
    "Since uflexltd.com runs on PHP with the same Apache server, the implementation approach is identical."
)

add_divider()

# ── B1-B5 ──
doc.add_heading("B1 - B5. Security Headers", level=2)

p = doc.add_paragraph(
    "All five security headers (CSP, X-Content-Type-Options, Referrer-Policy, Permissions-Policy, "
    "Cache-Control) can be applied to the production site using the same .htaccess approach described "
    "in Section A."
)

add_info_box(
    "Important: The Content Security Policy for the production site may need to be adjusted "
    "to include any additional third-party scripts or resources that the PHP site loads "
    "(e.g., different analytics, embedded videos, social media widgets). Before applying, "
    "we recommend reviewing which external resources the production site uses."
)

p = doc.add_paragraph(
    "The .htaccess block from Section A can be added to the production site's public_html/ directory. "
    "The CSP script-src and style-src directives should be reviewed and updated to include "
    "any production-specific third-party domains."
)

add_divider()

# ── B6 ──
doc.add_heading("B6. Infrastructure Hardening", level=2)

p = doc.add_paragraph(
    "Since both sites share the same server, the following improvements apply equally to both:"
)

add_styled_table(
    ["Improvement", "Applies To", "Notes"],
    [
        ["Hide Server Information", "Both sites", "Single .htaccess change per site"],
        ["Web Application Firewall (Cloudflare)", "Both sites", "DNS-level, protects all subdomains"],
        ["Rate Limiting on Forms", "Both sites", "Implement in respective form handlers"],
        ["Restrict Admin Port Access", "Server level", "Contact hosting provider"],
    ],
    col_widths=[2.2, 1.5, 2.5]
)

add_divider()

# ── B7 ──
doc.add_heading("B7. Privacy & Legal Pages", level=2)

p = doc.add_paragraph(
    "Both sites should include accessible privacy and legal pages. "
    "These were flagged in the assessment as informational disclosures. "
    "They should be created by the legal/content team and published on both domains:"
)

add_bullet("/privacy-policy/")
add_bullet("/terms-of-use/")
add_bullet("/cookie-policy/")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# SECTION C
# ═══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SECTION C")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = DARK_BLUE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Server-Level Updates")
run.font.size = Pt(18)
run.font.color.rgb = ACCENT_BLUE
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Applies to Both Sites (Shared Server)")
run.font.size = Pt(12)
run.font.color.rgb = MEDIUM_GRAY
run.font.name = 'Calibri'

doc.add_paragraph()

p = doc.add_paragraph(
    "These two items require action from the hosting provider or server administrator. "
    "They are not application-level changes and cannot be resolved through code deployment."
)

add_divider()

# ── C1: Exim ──
doc.add_heading("C1. Exim Mail Server Update", level=2)
add_severity_badge("HIGH (CVSS 7.8)", "high")

p = doc.add_paragraph()
run = p.add_run("Reference: ")
run.bold = True
run = p.add_run("CVE-2025-30232")

doc.add_heading("In Simple Terms", level=3)
p = doc.add_paragraph(
    "The server uses Exim (version 4.98.1) to handle email. A known security issue in this "
    "version could allow someone with server access to gain elevated privileges. "
    "This can be resolved by updating Exim to the latest patched version."
)

doc.add_heading("Recommended Action", level=3)
add_styled_table(
    ["Step", "Description", "Owner"],
    [
        ["1", "Update Exim to version 4.98.2 or later", "Hosting Provider"],
        ["2", "Verify the update was applied", "Server Admin"],
        ["3", "Restrict shell access to only authorized users", "Server Admin"],
        ["4", "Consider migrating to a managed email service (Google Workspace, Microsoft 365)", "IT Team"],
    ],
    col_widths=[0.5, 3.5, 2.2]
)

add_info_box(
    "For shared hosting: Submit a support ticket to the hosting provider "
    "requesting the Exim update, referencing CVE-2025-30232."
)

add_divider()

# ── C2: OpenSSH ──
doc.add_heading("C2. OpenSSH Update", level=2)
add_severity_badge("HIGH (CVSS 7.5 - 8.1)", "high")

p = doc.add_paragraph()
run = p.add_run("Reference: ")
run.bold = True
run = p.add_run("CVE-2026-35414")

doc.add_heading("In Simple Terms", level=3)
p = doc.add_paragraph(
    "The server uses OpenSSH (version 8.7) for secure remote access. A known issue in this "
    "version could allow unauthorized access under certain authentication configurations. "
    "Updating to the latest version resolves this."
)

doc.add_heading("Recommended Action", level=3)
add_styled_table(
    ["Step", "Description", "Owner"],
    [
        ["1", "Update OpenSSH to version 10.3 or later", "Hosting Provider"],
        ["2", "Verify the update was applied", "Server Admin"],
        ["3", "Restrict SSH access to specific authorized IP addresses", "Server Admin"],
        ["4", "Use key-based authentication instead of passwords", "Server Admin"],
    ],
    col_widths=[0.5, 3.5, 2.2]
)

add_info_box(
    "For shared hosting: Submit a support ticket to the hosting provider "
    "requesting the OpenSSH update, referencing CVE-2026-35414."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# REMEDIATION ROADMAP
# ═══════════════════════════════════════════════════════════════

doc.add_heading("Remediation Roadmap", level=1)

doc.add_heading("Phase 1 \u2014 Immediate (Within 48 Hours)", level=2)

add_styled_table(
    ["Item", "Action", "Owner", "Status"],
    [
        ["C1", "Update Exim mail server to 4.98.2+", "Hosting Provider", "Pending"],
        ["C2", "Update OpenSSH to 10.3+", "Hosting Provider", "Pending"],
    ],
    col_widths=[0.6, 2.8, 1.6, 1.2],
    header_color="C0392B"
)

doc.add_heading("Phase 2 \u2014 This Week (Within 7 Days)", level=2)

add_styled_table(
    ["Item", "Action", "Owner", "Status"],
    [
        ["A1-A5, B1-B5", "Deploy security headers via .htaccess", "Development Team", "Pending"],
        ["A6, B6", "Deploy Cloudflare WAF", "IT / DevOps", "Pending"],
        ["A6, B6", "Restrict admin port access", "Hosting Provider", "Pending"],
        ["A6, B6", "Hide server version information", "Development Team", "Pending"],
    ],
    col_widths=[1.0, 2.4, 1.6, 1.2],
    header_color="E67E22"
)

doc.add_heading("Phase 3 \u2014 Next Sprint (Within 14 Days)", level=2)

add_styled_table(
    ["Item", "Action", "Owner", "Status"],
    [
        ["A6, B6", "Add rate limiting to form endpoints", "Development Team", "Pending"],
        ["B7", "Publish privacy policy and legal pages", "Legal / Content Team", "Pending"],
        ["\u2014", "Schedule follow-up VAPT scan to verify fixes", "Security Team", "Pending"],
    ],
    col_widths=[1.0, 2.4, 1.6, 1.2],
    header_color="2E75B6"
)

add_divider()

# ═══════════════════════════════════════════════════════════════
# SUMMARY
# ═══════════════════════════════════════════════════════════════

doc.add_heading("Summary of Impact", level=1)

p = doc.add_paragraph("Once all the above improvements are implemented:")

impacts = [
    ("Security Headers: ", "All 5 application security headers will be in place, protecting against XSS, clickjacking, and data leakage"),
    ("Server Updates: ", "Server software will be up to date, closing the 2 known high-severity vulnerabilities"),
    ("WAF Protection: ", "A Web Application Firewall will filter malicious traffic before it reaches the server"),
    ("Form Protection: ", "Form endpoints will be protected against automated abuse"),
    ("Information Hiding: ", "Server information will no longer be exposed to external scanners"),
    ("Compliance: ", "Both the beta and production sites will meet industry-standard security benchmarks"),
]

for bold_part, text in impacts:
    add_bullet(text, bold_part)

doc.add_paragraph()
add_divider()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Report based on VAPT assessment dated April 13, 2026.")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = MEDIUM_GRAY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Recommended follow-up scan: After remediation is complete (target: May 2026)")
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = MEDIUM_GRAY
run.font.name = 'Calibri'


# ─── Save ───
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "UFlex-Security-Enhancement-Report.docx")
doc.save(output_path)
print(f"Report saved to: {output_path}")
